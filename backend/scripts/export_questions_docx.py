#!/usr/bin/env python3
"""
Export cached question-bank items to a .docx for offline expert (teacher) review.

Pulls from topic_anchors.question_bank, filters out cached fallback placeholders
("API Rate Limit Hit ..." / "System error fallback"), groups by topic, renders each
question with its answer key, and appends a per-question rating table so the document
doubles as the RQ4 expert-appraisal instrument.

Usage:
    python scripts/export_questions_docx.py \
        --subject "Bahasa Inggeris" --language English \
        --out evidence/review_bahasa_inggeris.docx
"""
import os
import sys
import argparse

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from dotenv import load_dotenv
load_dotenv(override=True)
from supabase import create_client

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def is_fallback(q):
    """True for cached error placeholders that must not be shown to a reviewer."""
    blob = " ".join([
        str(q.get("question") or ""),
        str(q.get("distractor_rationale") or ""),
        str(q.get("model_answer") or ""),
    ])
    return ("API Rate Limit Hit" in blob) or ("System error fallback" in blob)


def add_meta_line(doc, q):
    bits = []
    if q.get("question_type"):
        bits.append(q["question_type"])
    if q.get("kbat_level"):
        bits.append(f"KBAT: {q['kbat_level']}")
    marks = q.get("total_marks") or q.get("max_marks")
    if marks:
        bits.append(f"{marks} marks")
    if bits:
        p = doc.add_paragraph()
        r = p.add_run(" · ".join(str(b) for b in bits))
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_labeled(doc, label, text):
    if not text:
        return
    p = doc.add_paragraph()
    lr = p.add_run(f"{label}: ")
    lr.bold = True
    p.add_run(str(text))


def render_question(doc, n, q):
    doc.add_heading(f"Q{n}", level=3)
    add_meta_line(doc, q)

    # optional reading/listening context
    for key, label in (("passage", "Passage"), ("stimulus", "Stimulus"),
                        ("source_excerpt", "Source excerpt")):
        add_labeled(doc, label, q.get(key))
    if q.get("audio_url"):
        add_labeled(doc, "Audio", q["audio_url"])

    # stem
    if q.get("question"):
        p = doc.add_paragraph()
        p.add_run(str(q["question"]))

    qt = q.get("question_type")

    if qt == "mcq":
        opts = q.get("options") or []
        correct = q.get("correct_answer")
        for opt in opts:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(opt))
            if str(opt) == str(correct):
                run.bold = True
                run.add_text("   ✓ (key)")
        add_labeled(doc, "Correct answer", correct)
        dr = q.get("distractor_rationale")
        if isinstance(dr, dict) and dr:
            add_labeled(doc, "Distractor rationale", "; ".join(f"{k}: {v}" for k, v in dr.items()))
        add_labeled(doc, "Notes", q.get("illustrative_notes"))

    elif qt == "short_answer":
        for sp in (q.get("sub_parts") or []):
            p = doc.add_paragraph()
            p.add_run(f"{sp.get('label','')} ({sp.get('marks','?')} marks) ").bold = True
            p.add_run(str(sp.get("question", "")))
            add_labeled(doc, "  Sample answer", sp.get("sample_answer"))

    elif qt == "essay":
        themes = q.get("themes")
        if themes:
            add_labeled(doc, "Themes", ", ".join(themes) if isinstance(themes, list) else themes)
        add_labeled(doc, "Model answer", q.get("model_answer"))
        bands = q.get("marking_rubric_bands")
        if bands:
            add_labeled(doc, "Marking rubric", bands if isinstance(bands, str) else str(bands))

    else:  # generic fallback for any other type
        for k in ("correct_answer", "sample_answer", "model_answer", "explanation"):
            add_labeled(doc, k.replace("_", " ").title(), q.get(k))

    add_rating_table(doc)


CRITERIA = [
    "Curriculum alignment (KSSM)",
    "Accuracy of answer key",
    "Clarity / wording",
    "Difficulty appropriate for level",
    "Usefulness for teaching",
]


def add_rating_table(doc):
    p = doc.add_paragraph()
    p.add_run("Teacher review — rate 1 (poor) to 5 (excellent):").bold = True
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Criterion"
    for i, v in enumerate("12345", start=1):
        hdr[i].text = v
    for c in CRITERIA:
        row = table.add_row().cells
        row[0].text = c
    cp = doc.add_paragraph()
    cp.add_run("Comments / corrections: ").bold = True
    cp.add_run("_" * 90)
    doc.add_paragraph("")  # spacer


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--subject", default="Bahasa Inggeris")
    ap.add_argument("--language", default="English")
    ap.add_argument("--out", default=None)
    args = ap.parse_args()

    out = args.out or f"evidence/review_{args.subject.lower().replace(' ', '_')}.docx"

    sb = create_client(os.getenv("SUPABASE_URL"), os.getenv("SUPABASE_KEY"))
    res = (sb.table("topic_anchors")
           .select("topic, subject, language, question_bank")
           .eq("subject", args.subject).eq("language", args.language).execute())
    rows = sorted(res.data or [], key=lambda r: (r.get("topic") or ""))

    doc = Document()
    doc.add_heading(f"{args.subject} — Question Bank for Expert Review", level=0)
    intro = doc.add_paragraph()
    intro.add_run(
        f"KSSM {args.subject} ({args.language} language). Below are the AI-generated questions "
        "currently cached in the platform's question bank, exported for offline review. For each "
        "question, please rate the criteria and note any corrections. Fallback/error placeholders "
        "have been excluded.")

    total = kept = dropped = 0
    n = 0
    for r in rows:
        qs = [q for q in (r.get("question_bank") or [])]
        total += len(qs)
        good = [q for q in qs if not is_fallback(q)]
        dropped += len(qs) - len(good)
        kept += len(good)
        if not good:
            continue
        doc.add_heading(f"Topic: {r.get('topic')}", level=1)
        for q in good:
            n += 1
            render_question(doc, n, q)

    # summary footer
    doc.add_page_break()
    doc.add_heading("Export summary", level=1)
    doc.add_paragraph(f"Subject: {args.subject}  |  Language: {args.language}")
    doc.add_paragraph(f"Reviewable questions exported: {kept}")
    doc.add_paragraph(f"Fallback/error placeholders excluded: {dropped}")
    doc.add_paragraph(f"Total items in bank (before filtering): {total}")

    doc.save(out)
    print(f"Wrote {out}")
    print(f"  reviewable={kept}  excluded_fallback={dropped}  total={total}")


if __name__ == "__main__":
    main()
